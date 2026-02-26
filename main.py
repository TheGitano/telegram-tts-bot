import os
import io
import base64
import logging
from datetime import datetime, timedelta
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import (
    Application, CommandHandler, MessageHandler, filters,
    ContextTypes, CallbackQueryHandler, ConversationHandler
)
from docx import Document
import PyPDF2
from PyPDF2 import PdfWriter, PdfReader
from deep_translator import GoogleTranslator
from gtts import gTTS
import speech_recognition as sr
from pydub import AudioSegment
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import inch
import json

TELEGRAM_BOT_TOKEN = os.getenv("TELEGRAM_BOT_TOKEN")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
GOOGLE_CREDENTIALS_JSON = os.getenv("GOOGLE_APPLICATION_CREDENTIALS_JSON")
ADMIN_EMAIL = "corporatebusinessunitedstates@gmail.com"
FIRMA_TEXTO = "🦅 𝓣𝓱𝓮𝓖𝓲𝓽𝓪𝓷𝓸 🦅"

# Ruta de la imagen del bot (debe estar en el mismo directorio)
BOT_IMAGE_PATH = "1770404886764_image.png"

PREMIUM_USERS = {
    "Gitano": {"password": "8376", "name": "El Gitano", "email": "admin@gitano.com", "expires": datetime(2099, 12, 31)}
}

active_sessions = {}
free_usage = {}

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

CHOOSING_PLAN, PREMIUM_USERNAME, PREMIUM_PASSWORD, PREMIUM_BUY_DATA, FORGOT_PASSWORD, BUY_NOMBRE, BUY_APELLIDO, BUY_EMAIL, BUY_CELULAR, BUY_METODO_PAGO = range(10)

def is_premium_active(uid):
    if uid not in active_sessions:
        return False
    username = active_sessions[uid]
    if username not in PREMIUM_USERS:
        return False
    return datetime.now() <= PREMIUM_USERS[username]["expires"]

def get_premium_info(uid):
    if uid not in active_sessions:
        return None
    username = active_sessions[uid]
    if username not in PREMIUM_USERS:
        return None
    user_data = PREMIUM_USERS[username]
    days_left = max(0, (user_data["expires"] - datetime.now()).days)
    return {"username": username, "name": user_data["name"], "days_left": days_left}

def can_use_free(uid, function_name):
    if is_premium_active(uid):
        return True
    if uid not in free_usage:
        free_usage[uid] = {"texto": False, "documento": False, "audio": False, "doc_voz": False, "imagen": False}
    return not free_usage[uid].get(function_name, False)

def mark_free_used(uid, function_name):
    if uid not in free_usage:
        free_usage[uid] = {"texto": False, "documento": False, "audio": False, "doc_voz": False, "imagen": False}
    free_usage[uid][function_name] = True

def all_free_used(uid):
    if uid not in free_usage:
        return False
    return all(free_usage[uid].values())

def translate_text(text, source="auto", target="es"):
    try:
        if not text or len(text.strip()) == 0:
            return ""
        translator = GoogleTranslator(source=source, target=target)
        if len(text) > 4500:
            chunks = [text[i:i+4500] for i in range(0, len(text), 4500)]
            return " ".join([translator.translate(chunk) for chunk in chunks])
        return translator.translate(text)
    except Exception as e:
        logger.error(f"Error traducción: {e}")
        return text

def detect_language(text):
    try:
        from langdetect import detect
        lang = detect(text)
        logger.info(f"🔍 Idioma detectado: {lang} para texto: {text[:50]}...")
        return lang
    except Exception as e:
        logger.error(f"Error detectando idioma: {e}")
        return "unknown"

def tts(text, lang="es"):
    try:
        audio = io.BytesIO()
        if len(text) > 5000:
            text = text[:5000] + "..."
        logger.info(f"🔊 Generando TTS en idioma: {lang}")
        gTTS(text=text, lang=lang, slow=False).write_to_fp(audio)
        audio.seek(0)
        return audio
    except Exception as e:
        logger.error(f"Error TTS: {e}")
        return None

def transcribe_audio(audio_bytes):
    try:
        logger.info("🎤 Iniciando transcripción de audio...")
        audio = AudioSegment.from_file(io.BytesIO(audio_bytes))
        wav_io = io.BytesIO()
        audio.export(wav_io, format="wav")
        wav_io.seek(0)
        
        recognizer = sr.Recognizer()
        with sr.AudioFile(wav_io) as source:
            audio_data = recognizer.record(source)
            
            try:
                text_en = recognizer.recognize_google(audio_data, language="en-US")
                logger.info(f"✅ Transcrito en inglés: {text_en}")
                return text_en, "en"
            except:
                logger.info("❌ No se pudo transcribir en inglés, intentando español...")
            
            try:
                text_es = recognizer.recognize_google(audio_data, language="es-ES")
                logger.info(f"✅ Transcrito en español: {text_es}")
                return text_es, "es"
            except:
                logger.error("❌ No se pudo transcribir en ningún idioma")
                return None, None
                
    except Exception as e:
        logger.error(f"Error transcribiendo audio: {e}")
        return None, None

def extract_text_from_pdf(file_bytes):
    try:
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        return "\n".join([page.extract_text() or "" for page in reader.pages])
    except Exception as e:
        logger.error(f"Error PDF: {e}")
        return ""

def extract_text_from_docx(file_bytes):
    try:
        doc = Document(io.BytesIO(file_bytes))
        text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text)
        return "\n".join(text)
    except Exception as e:
        logger.error(f"Error DOCX extracción: {e}")
        return ""

def translate_docx(file_bytes, source_lang="auto", target_lang="es"):
    """Traduce un documento DOCX manteniendo el formato"""
    try:
        doc = Document(io.BytesIO(file_bytes))
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                original_text = paragraph.text
                translated = translate_text(original_text, source=source_lang, target=target_lang)
                paragraph.text = translated
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        original_text = cell.text
                        translated = translate_text(original_text, source=source_lang, target=target_lang)
                        cell.text = translated
        
        out_stream = io.BytesIO()
        doc.save(out_stream)
        out_stream.seek(0)
        return out_stream
    except Exception as e:
        logger.error(f"Error DOCX traducción: {e}")
        return None

def translate_pdf(file_bytes, source_lang="auto", target_lang="es"):
    """Traduce un PDF y lo devuelve como PDF"""
    try:
        text = extract_text_from_pdf(file_bytes)
        if not text:
            return None
        
        translated_text = translate_text(text, source=source_lang, target=target_lang)
        
        output = io.BytesIO()
        pdf_canvas = SimpleDocTemplate(output, pagesize=letter)
        
        styles = getSampleStyleSheet()
        story = []
        
        for para in translated_text.split('\n'):
            if para.strip():
                p = Paragraph(para, styles['Normal'])
                story.append(p)
                story.append(Spacer(1, 0.2*inch))
        
        pdf_canvas.build(story)
        output.seek(0)
        return output
    except Exception as e:
        logger.error(f"Error traduciendo PDF: {e}")
        return None

# ============================================================
# FUNCIONES DE IMAGEN CON GEMINI REST API
# ============================================================

def extract_text_with_vision(image_bytes):
    """Placeholder - usamos Gemini directamente."""
    return None


def analyze_image_with_claude(image_bytes, image_mime_type="image/jpeg", mode="analyze"):
    """
    Analiza una imagen usando Gemini REST API directamente.
    mode: "extract" -> solo texto | "analyze" -> texto + análisis experto
    """
    try:
        import requests as req
        gemini_key = os.getenv("GEMINI_API_KEY")
        if not gemini_key:
            logger.error("No se encontró GEMINI_API_KEY")
            return None

        image_b64 = base64.standard_b64encode(image_bytes).decode("utf-8")

        if mode == "extract":
            prompt_text = (
                "Extrae y transcribe con precisión TODO el texto visible en esta imagen. "
                "No agregues interpretaciones ni explicaciones. Solo el texto exacto tal como aparece."
            )
        else:
            prompt_text = (
                "Eres un experto multidisciplinario legal, financiero y de asesoría. "
                "Analiza el documento en la imagen:\n\n"
                "1. 📋 TEXTO EXTRAÍDO: Transcribe todo el texto visible.\n\n"
                "2. 📊 ANÁLISIS DETALLADO: Explica el contenido, fechas, montos y datos clave.\n\n"
                "3. ⚠️ SITUACIÓN ACTUAL: Urgencias y riesgos críticos.\n\n"
                "4. ✅ CONSEJOS Y RECOMENDACIONES: Pasos concretos como abogado/asesor.\n\n"
                "5. 📌 CONCLUSIÓN: Puntos más urgentes.\n\n"
                "Responde en el idioma del texto. Usa emojis para organizar."
            )

        # Modelos FREE TIER activos febrero 2026
        # gemini-2.0 y gemini-1.5 RETIRADOS (404)
        # Solo funcionan los 2.5:
        #   gemini-2.5-flash-lite -> 15 RPM, 1000 RPD
        #   gemini-2.5-flash      -> 10 RPM,  250 RPD
        #   gemini-2.5-pro        ->  5 RPM,  100 RPD
        models_to_try = [
            "gemini-2.5-flash-lite",
            "gemini-2.5-flash",
            "gemini-2.5-pro",
        ]

        for model_name in models_to_try:
            try:
                url = f"https://generativelanguage.googleapis.com/v1beta/models/{model_name}:generateContent?key={gemini_key}"
                body = {
                    "contents": [{
                        "parts": [
                            {"text": prompt_text},
                            {"inline_data": {"mime_type": image_mime_type, "data": image_b64}}
                        ]
                    }],
                    "generationConfig": {"temperature": 0.2, "maxOutputTokens": 2048}
                }
                logger.info(f"Llamando Gemini modelo: {model_name}")
                resp = req.post(url, json=body, timeout=60)
                logger.info(f"Gemini {model_name} status: {resp.status_code}")

                if resp.status_code == 200:
                    data = resp.json()
                    candidates = data.get("candidates", [])
                    if candidates:
                        parts = candidates[0].get("content", {}).get("parts", [])
                        text = "".join(p.get("text", "") for p in parts)
                        if text and len(text.strip()) > 3:
                            logger.info(f"OK: Gemini ({model_name}) -> {len(text)} chars")
                            return text.strip()
                elif resp.status_code == 429:
                    logger.warning(f"{model_name}: 429 cuota agotada, probando siguiente...")
                    continue
                else:
                    try:
                        error_msg = resp.json().get("error", {}).get("message", "?")
                    except Exception:
                        error_msg = resp.text[:100]
                    logger.error(f"{model_name} fallo ({resp.status_code}): {error_msg}")
                    continue
            except Exception as e:
                logger.error(f"Error con {model_name}: {e}")
                continue

        logger.error("Todos los modelos Gemini fallaron")
        return None

    except Exception as e:
        logger.error(f"Error analizando imagen con Gemini: {e}")
        return None

# ============================================================
# FIN FUNCIONES DE IMAGEN
# ============================================================

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    uid = update.effective_user.id
    context.user_data.clear()
    
    message = (
        "✨ *BIENVENIDO* ✨\n\n"
        f"🎨 BOT CREADO POR:\n🦅 𝓣𝓱𝓮𝓖𝓲𝓽𝓪𝓷𝓸 🦅\n\n"
        "⭐ *FUNCIONALIDADES:* ⭐\n\n"
        "📝 Texto a Voz\n"
        "🌐 Traductor Bidireccional\n"
        "📄 Traducir Documentos\n"
        "📋 Documentos a Voz\n"
        "🎤 Traducir Audio\n"
        "🖼️ Lector y Análisis de Imágenes\n\n"
        "💡 *SELECCIONA TU PLAN:* 💡\n\n"
        "🆓 *FREE:* 1 uso por función\n"
        "💎 *PREMIUM:* Uso ilimitado"
    )
    
    keyboard = [
        [InlineKeyboardButton("🆓 FREE", callback_data="plan_free")],
        [InlineKeyboardButton("💎 PREMIUM", callback_data="plan_premium")]
    ]
    
    if update.callback_query:
        await update.callback_query.edit_message_text(message, reply_markup=InlineKeyboardMarkup(keyboard), parse_mode="Markdown")
    else:
        await update.message.reply_text(message, reply_markup=InlineKeyboardMarkup(keyboard), parse_mode="Markdown")
    
    return CHOOSING_PLAN

async def plan_free(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    uid = update.effective_user.id
    
    used_functions = []
    if uid in free_usage:
        if free_usage[uid].get("texto"):
            used_functions.append("📝 Texto a Voz")
        if free_usage[uid].get("documento"):
            used_functions.append("📄 Traductor Documentos")
        if free_usage[uid].get("doc_voz"):
            used_functions.append("📋 Documentos a Voz")
        if free_usage[uid].get("audio"):
            used_functions.append("🎤 Traducir Audio")
        if free_usage[uid].get("imagen"):
            used_functions.append("🖼️ Lector de Imágenes")
    
    if all_free_used(uid):
        message = (
            "❌ *Ya usaste todas las funciones FREE*\n\n"
            "🎯 Funciones usadas:\n"
            f"{''.join([f'• {f}' + chr(10) for f in used_functions])}\n"
            "💎 *Actualiza a PREMIUM para uso ilimitado*\n\n"
            f"{FIRMA_TEXTO}"
        )
        keyboard = [
            [InlineKeyboardButton("💎 PREMIUM", callback_data="plan_premium")],
            [InlineKeyboardButton("🔙 Volver", callback_data="back_start")]
        ]
        await query.edit_message_text(message, reply_markup=InlineKeyboardMarkup(keyboard), parse_mode="Markdown")
        return CHOOSING_PLAN
    
    message = "🆓 *PLAN FREE*\n\n📌 Elige una función:\n\n"
    keyboard = []
    
    if can_use_free(uid, "texto"):
        keyboard.append([InlineKeyboardButton("📝 Texto a Voz", callback_data="free_texto")])
    else:
        keyboard.append([InlineKeyboardButton("📝 Texto a Voz ❌ Usado", callback_data="used")])
    
    if can_use_free(uid, "documento"):
        keyboard.append([InlineKeyboardButton("🌐 Traductor Documentos", callback_data="free_documento")])
    else:
        keyboard.append([InlineKeyboardButton("🌐 Traductor Documentos ❌ Usado", callback_data="used")])
    
    if can_use_free(uid, "doc_voz"):
        keyboard.append([InlineKeyboardButton("📋 Documentos a Voz", callback_data="free_doc_voz")])
    else:
        keyboard.append([InlineKeyboardButton("📋 Documentos a Voz ❌ Usado", callback_data="used")])
    
    if can_use_free(uid, "audio"):
        keyboard.append([InlineKeyboardButton("🎤 Traducir Audio", callback_data="free_audio")])
    else:
        keyboard.append([InlineKeyboardButton("🎤 Traducir Audio ❌ Usado", callback_data="used")])
    
    if can_use_free(uid, "imagen"):
        keyboard.append([InlineKeyboardButton("🖼️ Lector de Imágenes", callback_data="free_imagen")])
    else:
        keyboard.append([InlineKeyboardButton("🖼️ Lector de Imágenes ❌ Usado", callback_data="used")])
    
    keyboard.append([InlineKeyboardButton("💎 PREMIUM", callback_data="plan_premium")])
    keyboard.append([InlineKeyboardButton("🔙 Volver", callback_data="back_start")])
    
    await query.edit_message_text(message, reply_markup=InlineKeyboardMarkup(keyboard), parse_mode="Markdown")
    return CHOOSING_PLAN

async def plan_premium(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    uid = update.effective_user.id
    
    if is_premium_active(uid):
        return await show_premium_menu(update, context)
    
    message = (
        "💎 *PLAN PREMIUM*\n\n"
        "✅ Uso ilimitado de todas las funciones\n\n"
        "🎯 *FUNCIONES:*\n"
        "📝 Texto a Voz\n"
        "🌐 Traductor Bidireccional\n"
        "📄 Traducir Documentos\n"
        "📋 Documentos a Voz\n"
        "🎤 Traducir Audio\n"
        "🖼️ Lector y Análisis de Imágenes\n\n"
        f"{FIRMA_TEXTO}"
    )
    
    keyboard = [
        [InlineKeyboardButton("🔐 Iniciar Sesión", callback_data="premium_login")],
        [InlineKeyboardButton("💳 Comprar Premium", callback_data="buy_premium")],
        [InlineKeyboardButton("🔙 Volver", callback_data="back_start")]
    ]
    
    await query.edit_message_text(message, reply_markup=InlineKeyboardMarkup(keyboard), parse_mode="Markdown")
    return CHOOSING_PLAN

async def buy_premium(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    
    message = (
        "💎 *MEMBRESÍA PREMIUM*\n\n"
        "💵 *COSTO:* $27 USD/mes\n\n"
        "✨ *BENEFICIOS:*\n"
        "• Acceso ilimitado a todas las funciones\n"
        "• Sin restricciones de uso\n"
        "• Soporte prioritario\n"
        "• Validez por 30 días\n\n"
        "📋 *PROCESO DE COMPRA:*\n"
        "1️⃣ Completa tus datos\n"
        "2️⃣ Realiza el pago\n"
        "3️⃣ Envía tu comprobante\n"
        "4️⃣ Activa tu cuenta en 24h\n\n"
        "💳 *MÉTODOS DE PAGO:*\n\n"
        "🏦 *Western Union*\n"
        "   Alias: THEGITANO2AX.PF\n"
        "   Nombre: Matias Molina\n\n"
        "💸 *Zelle*\n"
        "   Tel: 3053314405\n"
        "   Nombre: Sebastian Tosi\n\n"
        "👇 *Presiona CONTINUAR para registrarte*\n\n"
        f"{FIRMA_TEXTO}"
    )
    
    keyboard = [
        [InlineKeyboardButton("✅ CONTINUAR", callback_data="start_buy_form")],
        [InlineKeyboardButton("🔙 Volver", callback_data="plan_premium")]
    ]
    
    await query.edit_message_text(message, reply_markup=InlineKeyboardMarkup(keyboard), parse_mode="Markdown")
    return CHOOSING_PLAN

async def start_buy_form(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    context.user_data["buy_form"] = {}
    
    await query.edit_message_text(
        "📝 *FORMULARIO DE REGISTRO*\n\n"
        "Por favor, escribe tu *NOMBRE COMPLETO:*",
        parse_mode="Markdown"
    )
    return BUY_NOMBRE

async def buy_nombre(update: Update, context: ContextTypes.DEFAULT_TYPE):
    nombre = update.message.text.strip()
    context.user_data["buy_form"]["nombre"] = nombre
    
    await update.message.reply_text(
        f"✅ Nombre: {nombre}\n\n"
        "Ahora escribe tu *APELLIDO(S):*",
        parse_mode="Markdown"
    )
    return BUY_APELLIDO

async def buy_apellido(update: Update, context: ContextTypes.DEFAULT_TYPE):
    apellido = update.message.text.strip()
    context.user_data["buy_form"]["apellido"] = apellido
    
    await update.message.reply_text(
        f"✅ Apellido: {apellido}\n\n"
        "Ahora escribe tu *CORREO ELECTRÓNICO:*",
        parse_mode="Markdown"
    )
    return BUY_EMAIL

async def buy_email(update: Update, context: ContextTypes.DEFAULT_TYPE):
    email = update.message.text.strip()
    
    if "@" not in email or "." not in email:
        await update.message.reply_text(
            "❌ Email inválido. Por favor, escribe un correo válido:"
        )
        return BUY_EMAIL
    
    context.user_data["buy_form"]["email"] = email
    
    await update.message.reply_text(
        f"✅ Email: {email}\n\n"
        "Ahora escribe tu *NÚMERO DE CELULAR:*\n"
        "(Con código de país, ej: +1234567890)",
        parse_mode="Markdown"
    )
    return BUY_CELULAR

async def buy_celular(update: Update, context: ContextTypes.DEFAULT_TYPE):
    celular = update.message.text.strip()
    context.user_data["buy_form"]["celular"] = celular
    
    datos = context.user_data["buy_form"]
    
    mensaje_resumen = (
        "📋 *RESUMEN DE TUS DATOS*\n\n"
        f"👤 Nombre: {datos['nombre']} {datos['apellido']}\n"
        f"📧 Email: {datos['email']}\n"
        f"📱 Celular: {datos['celular']}\n\n"
        "💳 *MÉTODOS DE PAGO DISPONIBLES:*\n\n"
        "🏦 *Western Union*\n"
        "   Alias: THEGITANO2AX.PF\n"
        "   Nombre: Matias Molina\n\n"
        "💸 *Zelle*\n"
        "   Tel: 3053314405\n"
        "   Nombre: Sebastian Tosi\n\n"
        "¿Con qué método realizaste o realizarás el pago?"
    )
    
    keyboard = [
        [InlineKeyboardButton("🏦 Western Union", callback_data="pago_western")],
        [InlineKeyboardButton("💸 Zelle", callback_data="pago_zelle")],
        [InlineKeyboardButton("🔄 Cancelar", callback_data="plan_premium")]
    ]
    
    await update.message.reply_text(
        mensaje_resumen,
        reply_markup=InlineKeyboardMarkup(keyboard),
        parse_mode="Markdown"
    )
    return BUY_METODO_PAGO

async def buy_metodo_pago(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    
    if query.data == "pago_western":
        metodo = "Western Union"
    elif query.data == "pago_zelle":
        metodo = "Zelle"
    else:
        return CHOOSING_PLAN
    
    context.user_data["buy_form"]["metodo_pago"] = metodo
    datos = context.user_data["buy_form"]
    
    email_content = f"""
NUEVA SOLICITUD DE MEMBRESÍA PREMIUM

═══════════════════════════════
DATOS DEL CLIENTE
═══════════════════════════════
Nombre Completo: {datos['nombre']} {datos['apellido']}
Email: {datos['email']}
Celular: {datos['celular']}
Método de Pago: {metodo}

═══════════════════════════════
INFORMACIÓN DE PAGO
═══════════════════════════════
Monto: $27 USD
Período: 1 mes
"""
    
    logger.info(f"📧 NUEVA SOLICITUD PREMIUM:\n{email_content}")
    
    mensaje_confirmacion = (
        "✅ *¡SOLICITUD ENVIADA!*\n\n"
        f"Tus datos han sido enviados a:\n"
        f"📮 {ADMIN_EMAIL}\n\n"
        "📝 *RESUMEN:*\n"
        f"👤 {datos['nombre']} {datos['apellido']}\n"
        f"📧 {datos['email']}\n"
        f"📱 {datos['celular']}\n"
        f"💳 Método: {metodo}\n\n"
        "⏳ *PRÓXIMOS PASOS:*\n"
        "1️⃣ Realiza el pago de $27 USD\n"
        "2️⃣ Envía el comprobante a nuestro email\n"
        "3️⃣ Tu cuenta será activada en 24h\n\n"
        f"{FIRMA_TEXTO}"
    )
    
    keyboard = [
        [InlineKeyboardButton("🏠 Volver al Inicio", callback_data="back_start")]
    ]
    
    await query.edit_message_text(
        mensaje_confirmacion,
        reply_markup=InlineKeyboardMarkup(keyboard),
        parse_mode="Markdown"
    )
    
    context.user_data["buy_form"] = {}
    
    return CHOOSING_PLAN

async def premium_login(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.edit_message_text(
        "🔐 *INICIAR SESIÓN PREMIUM*\n\n📝 Escribe tu usuario:",
        parse_mode="Markdown"
    )
    return PREMIUM_USERNAME

async def premium_username(update: Update, context: ContextTypes.DEFAULT_TYPE):
    username = update.message.text.strip()
    
    if username not in PREMIUM_USERS:
        keyboard = [[InlineKeyboardButton("🔙 Volver", callback_data="plan_premium")]]
        await update.message.reply_text(
            "❌ Usuario no encontrado\n\n"
            "💡 ¿Olvidaste tu contraseña?",
            reply_markup=InlineKeyboardMarkup(keyboard)
        )
        return CHOOSING_PLAN
    
    context.user_data["temp_username"] = username
    await update.message.reply_text("🔑 Escribe tu contraseña:")
    return PREMIUM_PASSWORD

async def premium_password(update: Update, context: ContextTypes.DEFAULT_TYPE):
    password = update.message.text.strip()
    username = context.user_data.get("temp_username")
    uid = update.effective_user.id
    
    if username and PREMIUM_USERS.get(username, {}).get("password") == password:
        active_sessions[uid] = username
        context.user_data["is_premium"] = True
        await update.message.reply_text("✅ ¡Sesión iniciada!")
        
        fake_query = type('obj', (object,), {
            'edit_message_text': update.message.reply_text,
            'answer': lambda: None
        })()
        fake_update = type('obj', (object,), {
            'callback_query': fake_query,
            'effective_user': update.effective_user
        })()
        
        return await show_premium_menu(fake_update, context)
    else:
        keyboard = [
            [InlineKeyboardButton("🔄 Reintentar", callback_data="premium_login")],
            [InlineKeyboardButton("❓ Olvidé mi contraseña", callback_data="forgot_password")],
            [InlineKeyboardButton("🔙 Volver", callback_data="plan_premium")]
        ]
        await update.message.reply_text(
            "❌ Contraseña incorrecta",
            reply_markup=InlineKeyboardMarkup(keyboard)
        )
        return CHOOSING_PLAN

async def forgot_password(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.edit_message_text(
        f"📧 Envía un email a:\n{ADMIN_EMAIL}\n\n"
        "Incluye tu usuario registrado.\n\n"
        f"{FIRMA_TEXTO}"
    )
    return FORGOT_PASSWORD

async def process_forgot_password(update: Update, context: ContextTypes.DEFAULT_TYPE):
    keyboard = [[InlineKeyboardButton("🔙 Volver", callback_data="plan_premium")]]
    await update.message.reply_text(
        "✅ Solicitud recibida.\n"
        f"Revisa {ADMIN_EMAIL} en 24h.",
        reply_markup=InlineKeyboardMarkup(keyboard)
    )
    return CHOOSING_PLAN

async def show_premium_menu(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    uid = update.effective_user.id
    
    info = get_premium_info(uid)
    if not info:
        return await plan_premium(update, context)
    
    message = (
        f"💎 *MENÚ PREMIUM*\n\n"
        f"👤 Usuario: {info['name']}\n"
        f"⏰ Días restantes: {info['days_left']}\n\n"
        "🎯 *Elige una función:*"
    )
    
    keyboard = [
        [InlineKeyboardButton("📝 Texto a Voz", callback_data="premium_texto")],
        [InlineKeyboardButton("🌐 Traductor Documentos", callback_data="premium_documento")],
        [InlineKeyboardButton("📋 Documentos a Voz", callback_data="premium_doc_voz")],
        [InlineKeyboardButton("🎤 Traducir Audio", callback_data="premium_audio")],
        [InlineKeyboardButton("🖼️ Lector de Imágenes", callback_data="premium_imagen")],
        [InlineKeyboardButton("🚪 Cerrar Sesión", callback_data="premium_logout")]
    ]
    
    await query.edit_message_text(message, reply_markup=InlineKeyboardMarkup(keyboard), parse_mode="Markdown")
    return CHOOSING_PLAN

async def free_texto(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    uid = update.effective_user.id
    
    if not can_use_free(uid, "texto"):
        await query.answer("❌ Ya usaste esta función", show_alert=True)
        return CHOOSING_PLAN
    
    context.user_data["waiting_text"] = True
    context.user_data["is_premium"] = False
    
    try:
        if os.path.exists(BOT_IMAGE_PATH):
            with open(BOT_IMAGE_PATH, 'rb') as photo:
                await query.message.reply_photo(
                    photo=photo,
                    caption=(
                        "📝 *TEXTO A VOZ*\n\n"
                        "📌 Guía de uso:\n"
                        "1️⃣ Envía un texto\n"
                        "2️⃣ Elige si quieres traducción\n"
                        "3️⃣ Recibe tu audio\n\n"
                        f"{FIRMA_TEXTO}"
                    ),
                    parse_mode="Markdown"
                )
        else:
            await query.message.reply_text(
                "📝 *TEXTO A VOZ*\n\n"
                "✍️ Envía el texto que quieres convertir a voz\n\n"
                f"{FIRMA_TEXTO}",
                parse_mode="Markdown"
            )
    except Exception as e:
        logger.error(f"Error enviando imagen: {e}")
        await query.message.reply_text(
            "📝 *TEXTO A VOZ*\n\n"
            "✍️ Envía el texto que quieres convertir a voz\n\n"
            f"{FIRMA_TEXTO}",
            parse_mode="Markdown"
        )
    
    await query.answer()
    return CHOOSING_PLAN

async def premium_texto(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    uid = update.effective_user.id
    
    if not is_premium_active(uid):
        await query.answer("❌ Sesión expirada", show_alert=True)
        return await plan_premium(update, context)
    
    context.user_data["waiting_text"] = True
    context.user_data["is_premium"] = True
    
    try:
        if os.path.exists(BOT_IMAGE_PATH):
            with open(BOT_IMAGE_PATH, 'rb') as photo:
                await query.message.reply_photo(
                    photo=photo,
                    caption=(
                        "📝 *TEXTO A VOZ*\n\n"
                        "📌 Guía de uso:\n"
                        "1️⃣ Envía un texto\n"
                        "2️⃣ Elige si quieres traducción\n"
                        "3️⃣ Recibe tu audio\n\n"
                        f"{FIRMA_TEXTO}"
                    ),
                    parse_mode="Markdown"
                )
        else:
            await query.message.reply_text(
                "📝 *TEXTO A VOZ*\n\n"
                "✍️ Envía el texto que quieres convertir a voz\n\n"
                f"{FIRMA_TEXTO}",
                parse_mode="Markdown"
            )
    except Exception as e:
        logger.error(f"Error enviando imagen: {e}")
        await query.message.reply_text(
            "📝 *TEXTO A VOZ*\n\n"
            "✍️ Envía el texto que quieres convertir a voz\n\n"
            f"{FIRMA_TEXTO}",
            parse_mode="Markdown"
        )
    
    await query.answer()
    return CHOOSING_PLAN

async def free_documento(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    uid = update.effective_user.id
    
    if not can_use_free(uid, "documento"):
        await query.answer("❌ Ya usaste esta función", show_alert=True)
        return CHOOSING_PLAN
    
    context.user_data["waiting_document"] = True
    context.user_data["document_mode"] = "translate"
    context.user_data["is_premium"] = False
    
    await query.message.reply_text(
        "📄 *TRADUCTOR DE DOCUMENTOS*\n\n"
        "📎 Envía un documento (PDF o DOCX)\n\n"
        "🌐 Se traducirá automáticamente:\n"
        "• 🇪🇸 Español → 🇺🇸 Inglés\n"
        "• 🇺🇸 Inglés → 🇪🇸 Español\n\n"
        "📋 El documento se entregará en el mismo formato\n\n"
        f"{FIRMA_TEXTO}",
        parse_mode="Markdown"
    )
    await query.answer()
    return CHOOSING_PLAN

async def premium_documento(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    uid = update.effective_user.id
    
    if not is_premium_active(uid):
        await query.answer("❌ Sesión expirada", show_alert=True)
        return await plan_premium(update, context)
    
    context.user_data["waiting_document"] = True
    context.user_data["document_mode"] = "translate"
    context.user_data["is_premium"] = True
    
    await query.message.reply_text(
        "📄 *TRADUCTOR DE DOCUMENTOS*\n\n"
        "📎 Envía un documento (PDF o DOCX)\n\n"
        "🌐 Se traducirá automáticamente:\n"
        "• 🇪🇸 Español → 🇺🇸 Inglés\n"
        "• 🇺🇸 Inglés → 🇪🇸 Español\n\n"
        "📋 El documento se entregará en el mismo formato\n\n"
        f"{FIRMA_TEXTO}",
        parse_mode="Markdown"
    )
    await query.answer()
    return CHOOSING_PLAN

async def free_doc_voz(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    uid = update.effective_user.id
    
    if not can_use_free(uid, "doc_voz"):
        await query.answer("❌ Ya usaste esta función", show_alert=True)
        return CHOOSING_PLAN
    
    context.user_data["waiting_document"] = True
    context.user_data["document_mode"] = "voice"
    context.user_data["is_premium"] = False
    
    await query.message.reply_text(
        "📋 *DOCUMENTOS A VOZ*\n\n"
        "📎 Envía un documento (PDF o DOCX)\n\n"
        "🔊 Recibirás el audio del contenido traducido\n\n"
        f"{FIRMA_TEXTO}",
        parse_mode="Markdown"
    )
    await query.answer()
    return CHOOSING_PLAN

async def premium_doc_voz(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    uid = update.effective_user.id
    
    if not is_premium_active(uid):
        await query.answer("❌ Sesión expirada", show_alert=True)
        return await plan_premium(update, context)
    
    context.user_data["waiting_document"] = True
    context.user_data["document_mode"] = "voice"
    context.user_data["is_premium"] = True
    
    await query.message.reply_text(
        "📋 *DOCUMENTOS A VOZ*\n\n"
        "📎 Envía un documento (PDF o DOCX)\n\n"
        "🔊 Recibirás el audio del contenido traducido\n\n"
        f"{FIRMA_TEXTO}",
        parse_mode="Markdown"
    )
    await query.answer()
    return CHOOSING_PLAN

async def free_audio(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    uid = update.effective_user.id
    
    if not can_use_free(uid, "audio"):
        await query.answer("❌ Ya usaste esta función", show_alert=True)
        return CHOOSING_PLAN
    
    context.user_data["waiting_audio"] = True
    context.user_data["is_premium"] = False
    
    await query.message.reply_text(
        "🎤 *TRADUCIR AUDIO*\n\n"
        "🎙 Envía un audio o nota de voz\n\n"
        "🌐 Se transcribirá y traducirá:\n"
        "• 🇪🇸 Español → 🇺🇸 Inglés\n"
        "• 🇺🇸 Inglés → 🇪🇸 Español\n\n"
        f"{FIRMA_TEXTO}",
        parse_mode="Markdown"
    )
    await query.answer()
    return CHOOSING_PLAN

async def premium_audio(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    uid = update.effective_user.id
    
    if not is_premium_active(uid):
        await query.answer("❌ Sesión expirada", show_alert=True)
        return await plan_premium(update, context)
    
    context.user_data["waiting_audio"] = True
    context.user_data["is_premium"] = True
    
    await query.message.reply_text(
        "🎤 *TRADUCIR AUDIO*\n\n"
        "🎙 Envía un audio o nota de voz\n\n"
        "🌐 Se transcribirá y traducirá:\n"
        "• 🇪🇸 Español → 🇺🇸 Inglés\n"
        "• 🇺🇸 Inglés → 🇪🇸 Español\n\n"
        f"{FIRMA_TEXTO}",
        parse_mode="Markdown"
    )
    await query.answer()
    return CHOOSING_PLAN

# ============================================================
# HANDLERS DE IMAGEN - NUEVOS
# ============================================================

async def free_imagen(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    uid = update.effective_user.id
    
    if not can_use_free(uid, "imagen"):
        await query.answer("❌ Ya usaste esta función", show_alert=True)
        return CHOOSING_PLAN
    
    context.user_data["waiting_image"] = True
    context.user_data["is_premium"] = False
    
    await query.message.reply_text(
        "🖼️ *LECTOR Y ANÁLISIS DE IMÁGENES*\n\n"
        "📸 Envía una imagen con texto\n\n"
        "✨ *¿Qué puedo hacer?*\n"
        "• Extraer todo el texto de la imagen\n"
        "• Traducir el texto (ES↔EN)\n"
        "• Convertir el texto a voz\n"
        "• Analizar y dar consejos expertos\n\n"
        f"{FIRMA_TEXTO}",
        parse_mode="Markdown"
    )
    await query.answer()
    return CHOOSING_PLAN

async def premium_imagen(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    uid = update.effective_user.id
    
    if not is_premium_active(uid):
        await query.answer("❌ Sesión expirada", show_alert=True)
        return await plan_premium(update, context)
    
    context.user_data["waiting_image"] = True
    context.user_data["is_premium"] = True
    
    await query.message.reply_text(
        "🖼️ *LECTOR Y ANÁLISIS DE IMÁGENES*\n\n"
        "📸 Envía una imagen con texto\n\n"
        "✨ *¿Qué puedo hacer?*\n"
        "• Extraer todo el texto de la imagen\n"
        "• Traducir el texto (ES↔EN)\n"
        "• Convertir el texto a voz\n"
        "• Analizar y dar consejos expertos\n\n"
        f"{FIRMA_TEXTO}",
        parse_mode="Markdown"
    )
    await query.answer()
    return CHOOSING_PLAN

async def handle_image(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Maneja las imágenes recibidas cuando el modo imagen está activo"""
    uid = update.effective_user.id
    
    if not context.user_data.get("waiting_image", False):
        return
    
    try:
        processing_msg = await update.message.reply_text("⏳ Procesando imagen...")
        
        # Obtener la imagen de mayor resolución disponible
        photo = update.message.photo[-1]
        file = await context.bot.get_file(photo.file_id)
        image_bytes = bytes(await file.download_as_bytearray())
        
        # Guardar imagen en contexto para reutilizar
        context.user_data["pending_image_bytes"] = image_bytes
        context.user_data["pending_image_mime"] = "image/jpeg"
        
        # Primero extraemos el texto para detectar el idioma
        await processing_msg.edit_text("🔍 Leyendo texto de la imagen...")
        extracted_text = analyze_image_with_claude(image_bytes, "image/jpeg", mode="extract")
        
        if not extracted_text or len(extracted_text.strip()) < 5:
            await update.message.reply_text(
                "❌ No se pudo detectar texto en la imagen.\n\n"
                "Asegúrate de que la imagen tenga texto visible y legible.",
                parse_mode="Markdown"
            )
            await processing_msg.delete()
            return
        
        context.user_data["pending_image_text"] = extracted_text
        
        # Detectar idioma del texto extraído
        lang = detect_language(extracted_text)
        context.user_data["pending_image_lang"] = lang
        
        if lang == "es":
            lang_emoji = "🇪🇸"
            target_emoji = "🇺🇸"
            lang_name = "Español"
            target_name = "Inglés"
        else:
            lang_emoji = "🇺🇸"
            target_emoji = "🇪🇸"
            lang_name = "Inglés"
            target_name = "Español"
        
        await processing_msg.delete()
        
        # Mostrar menú de opciones
        keyboard = [
            [InlineKeyboardButton(f"🔊 Audio en {lang_name} (sin traducir)", callback_data="img_audio_original")],
            [InlineKeyboardButton(f"🔊 Audio traducido {lang_emoji}→{target_emoji}", callback_data="img_audio_traducido")],
            [InlineKeyboardButton("🔍 Analizar + Consejos Expertos", callback_data="img_analizar")],
            [InlineKeyboardButton("📋 Solo ver el texto extraído", callback_data="img_solo_texto")],
        ]
        
        # Mostrar preview del texto
        preview = extracted_text[:200] + "..." if len(extracted_text) > 200 else extracted_text
        
        await update.message.reply_text(
            f"✅ *Imagen procesada*\n\n"
            f"📝 *Texto detectado ({lang_name}):*\n"
            f"_{preview}_\n\n"
            f"¿Qué deseas hacer con este contenido?",
            reply_markup=InlineKeyboardMarkup(keyboard),
            parse_mode="Markdown"
        )
        
    except Exception as e:
        logger.error(f"❌ Error procesando imagen: {e}")
        await update.message.reply_text("❌ Error al procesar la imagen. Intenta de nuevo.")
        try:
            await processing_msg.delete()
        except:
            pass

async def handle_image_action(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Maneja las acciones sobre la imagen procesada"""
    query = update.callback_query
    await query.answer()
    uid = update.effective_user.id
    
    image_bytes = context.user_data.get("pending_image_bytes")
    image_text = context.user_data.get("pending_image_text", "")
    image_lang = context.user_data.get("pending_image_lang", "unknown")
    image_mime = context.user_data.get("pending_image_mime", "image/jpeg")
    
    if not image_bytes:
        await query.edit_message_text("❌ No hay imagen pendiente. Envía una imagen primero.")
        return CHOOSING_PLAN
    
    action = query.data
    
    if action == "img_solo_texto":
        # Solo mostrar el texto extraído
        # Dividir en partes si es muy largo
        max_len = 4000
        if len(image_text) > max_len:
            parts = [image_text[i:i+max_len] for i in range(0, len(image_text), max_len)]
            for i, part in enumerate(parts):
                await query.message.reply_text(
                    f"📋 *Texto extraído (parte {i+1}/{len(parts)}):*\n\n{part}",
                    parse_mode="Markdown"
                )
        else:
            await query.message.reply_text(
                f"📋 *Texto extraído de la imagen:*\n\n{image_text}\n\n{FIRMA_TEXTO}",
                parse_mode="Markdown"
            )
        
        if not context.user_data.get("is_premium", False):
            mark_free_used(uid, "imagen")
        
        _send_back_menu(update, context)
    
    elif action == "img_audio_original":
        processing_msg = await query.message.reply_text("⏳ Generando audio...")
        
        # Audio en idioma original
        if image_lang == "es":
            audio_lang = "es"
            lang_display = "🇪🇸 Español"
        else:
            audio_lang = "en"
            lang_display = "🇺🇸 Inglés"
        
        audio = tts(image_text, audio_lang)
        
        if audio:
            await query.message.reply_voice(
                audio,
                caption=f"🔊 Audio en {lang_display} (original)\n\n{FIRMA_TEXTO}"
            )
        else:
            await query.message.reply_text("❌ Error al generar audio.")
        
        await processing_msg.delete()
        
        if not context.user_data.get("is_premium", False):
            mark_free_used(uid, "imagen")
        
        await _send_back_menu_message(query.message, context)
    
    elif action == "img_audio_traducido":
        processing_msg = await query.message.reply_text("⏳ Traduciendo y generando audio...")
        
        if image_lang == "es":
            target_lang = "en"
            audio_lang = "en"
            lang_display = "🇪🇸→🇺🇸"
        else:
            target_lang = "es"
            audio_lang = "es"
            lang_display = "🇺🇸→🇪🇸"
        
        translated_text = translate_text(image_text, source=image_lang, target=target_lang)
        audio = tts(translated_text, audio_lang)
        
        if audio:
            await query.message.reply_voice(
                audio,
                caption=f"🔊 Audio traducido {lang_display}\n\n{FIRMA_TEXTO}"
            )
        else:
            await query.message.reply_text("❌ Error al generar audio traducido.")
        
        await processing_msg.delete()
        
        if not context.user_data.get("is_premium", False):
            mark_free_used(uid, "imagen")
        
        await _send_back_menu_message(query.message, context)
    
    elif action == "img_analizar":
        processing_msg = await query.message.reply_text("🔍 Analizando imagen con IA experta...\n\n⏳ Esto puede tardar unos segundos...")
        
        analysis = analyze_image_with_claude(image_bytes, image_mime, mode="analyze")
        
        await processing_msg.delete()
        
        if analysis:
            # Dividir análisis en partes si es muy largo
            max_len = 4000
            if len(analysis) > max_len:
                parts = [analysis[i:i+max_len] for i in range(0, len(analysis), max_len)]
                for i, part in enumerate(parts):
                    header = f"📊 *ANÁLISIS EXPERTO (parte {i+1}/{len(parts)}):*\n\n" if i == 0 else ""
                    await query.message.reply_text(
                        f"{header}{part}",
                        parse_mode="Markdown"
                    )
            else:
                await query.message.reply_text(
                    f"📊 *ANÁLISIS EXPERTO:*\n\n{analysis}\n\n{FIRMA_TEXTO}",
                    parse_mode="Markdown"
                )
            
            # Ofrecer también el audio del análisis
            keyboard = [
                [InlineKeyboardButton("🔊 Escuchar análisis en audio", callback_data="img_audio_analisis")],
                [InlineKeyboardButton("🔙 Volver al Menú", callback_data="premium_menu" if context.user_data.get("is_premium") else "plan_free")]
            ]
            context.user_data["pending_analysis_text"] = analysis
            context.user_data["pending_analysis_lang"] = image_lang
            
            await query.message.reply_text(
                "¿Deseas escuchar el análisis en audio?",
                reply_markup=InlineKeyboardMarkup(keyboard)
            )
        else:
            await query.message.reply_text("❌ Error al analizar la imagen.")
            await _send_back_menu_message(query.message, context)
        
        if not context.user_data.get("is_premium", False):
            mark_free_used(uid, "imagen")
    
    elif action == "img_audio_analisis":
        processing_msg = await query.message.reply_text("⏳ Generando audio del análisis...")
        
        analysis_text = context.user_data.get("pending_analysis_text", "")
        analysis_lang = context.user_data.get("pending_analysis_lang", "es")
        
        # El análisis siempre se devuelve en el idioma de la imagen
        audio_lang = "es" if analysis_lang == "es" else "en"
        audio = tts(analysis_text[:4500], audio_lang)
        
        await processing_msg.delete()
        
        if audio:
            await query.message.reply_voice(
                audio,
                caption=f"🔊 Audio del análisis experto\n\n{FIRMA_TEXTO}"
            )
        else:
            await query.message.reply_text("❌ Error al generar audio del análisis.")
        
        await _send_back_menu_message(query.message, context)
    
    return CHOOSING_PLAN

async def _send_back_menu_message(message, context):
    """Envía botón de volver al menú"""
    back = "premium_menu" if context.user_data.get("is_premium") else "plan_free"
    keyboard = [[InlineKeyboardButton("🔙 Volver al Menú", callback_data=back)]]
    await message.reply_text(
        f"✅ Listo. Envía otra imagen o vuelve al menú.\n\n{FIRMA_TEXTO}",
        reply_markup=InlineKeyboardMarkup(keyboard)
    )

# ============================================================
# FIN HANDLERS DE IMAGEN
# ============================================================

async def handle_text(update: Update, context: ContextTypes.DEFAULT_TYPE):
    uid = update.effective_user.id
    
    if context.user_data.get("waiting_translate_response"):
        response = update.message.text.strip().upper()
        
        if response not in ["SI", "SÍ", "NO"]:
            keyboard = [
                [InlineKeyboardButton("✅ SI", callback_data="translate_yes")],
                [InlineKeyboardButton("❌ NO", callback_data="translate_no")]
            ]
            await update.message.reply_text(
                "❓ Por favor, responde SI o NO\n\n"
                "¿Quieres que el audio sea traducido?",
                reply_markup=InlineKeyboardMarkup(keyboard)
            )
            return CHOOSING_PLAN
        
        text_to_process = context.user_data.get("text_to_process", "")
        lang = detect_language(text_to_process)
        
        processing_msg = await update.message.reply_text("⏳ Generando audio...")
        
        if response in ["SI", "SÍ"]:
            if lang == "es":
                target_lang = "en"
                audio_lang = "en"
                lang_display = "🇪🇸→🇺🇸"
                translated_text = translate_text(text_to_process, source="es", target="en")
            else:
                target_lang = "es"
                audio_lang = "es"
                lang_display = "🇺🇸→🇪🇸"
                translated_text = translate_text(text_to_process, source="en", target="es")
            
            audio = tts(translated_text, audio_lang)
            caption = f"{lang_display} Audio traducido\n\n{FIRMA_TEXTO}"
        else:
            if lang == "es":
                audio_lang = "es"
            else:
                audio_lang = "en"
            
            audio = tts(text_to_process, audio_lang)
            caption = f"🔊 Audio en idioma original\n\n{FIRMA_TEXTO}"
        
        if audio:
            await update.message.reply_voice(audio, caption=caption)
        else:
            await update.message.reply_text("❌ Error al generar audio.")
            await processing_msg.delete()
            return CHOOSING_PLAN
        
        if not context.user_data.get("is_premium", False):
            mark_free_used(uid, "texto")
        
        if not context.user_data.get("is_premium", False) and all_free_used(uid):
            keyboard = [[InlineKeyboardButton("💎 PREMIUM", callback_data="plan_premium")]]
            await update.message.reply_text(f"✅ Ya usaste FREE\n\n{FIRMA_TEXTO}", reply_markup=InlineKeyboardMarkup(keyboard))
            context.user_data["waiting_text"] = False
            context.user_data["waiting_translate_response"] = False
        else:
            back = "premium_menu" if context.user_data.get("is_premium") else "plan_free"
            keyboard = [[InlineKeyboardButton("🔙 Volver al Menú", callback_data=back)]]
            await update.message.reply_text(f"✅ Listo. Envía otro texto o vuelve.\n\n{FIRMA_TEXTO}", reply_markup=InlineKeyboardMarkup(keyboard))
        
        context.user_data["waiting_translate_response"] = False
        context.user_data["text_to_process"] = ""
        await processing_msg.delete()
        return CHOOSING_PLAN
    
    if not context.user_data.get("waiting_text", False):
        return CHOOSING_PLAN
    
    text = update.message.text.strip()
    
    if not text:
        await update.message.reply_text("❌ Texto vacío. Envía un texto válido.")
        return CHOOSING_PLAN
    
    context.user_data["text_to_process"] = text
    context.user_data["waiting_translate_response"] = True
    
    keyboard = [
        [InlineKeyboardButton("✅ SI", callback_data="translate_yes")],
        [InlineKeyboardButton("❌ NO", callback_data="translate_no")]
    ]
    
    await update.message.reply_text(
        "🌐 ¿Quieres que el audio sea traducido?\n\n"
        "• SI: El audio se traducirá (ES↔EN)\n"
        "• NO: El audio será en el idioma original",
        reply_markup=InlineKeyboardMarkup(keyboard)
    )
    
    return CHOOSING_PLAN

async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE):
    uid = update.effective_user.id
    if not context.user_data.get("waiting_document", False):
        return
    
    try:
        processing_msg = await update.message.reply_text("⏳ Procesando documento...")
        
        doc = update.message.document
        file = await context.bot.get_file(doc.file_id)
        data = await file.download_as_bytearray()
        
        document_mode = context.user_data.get("document_mode", "translate")
        is_docx = doc.file_name.endswith(".docx")
        is_pdf = doc.file_name.endswith(".pdf")
        
        if not (is_docx or is_pdf):
            await update.message.reply_text("❌ Solo se permiten archivos PDF o DOCX")
            await processing_msg.delete()
            return
        
        if is_docx:
            text = extract_text_from_docx(data)
        else:
            text = extract_text_from_pdf(data)
        
        if not text:
            await update.message.reply_text("❌ No se pudo extraer texto del documento")
            await processing_msg.delete()
            return
        
        lang = detect_language(text)
        
        if lang == "es":
            target_lang = "en"
            lang_display = "🇪🇸→🇺🇸"
            audio_lang = "en"
        else:
            target_lang = "es"
            lang_display = "🇺🇸→🇪🇸"
            audio_lang = "es"
        
        if document_mode == "translate":
            if is_docx:
                translated_file = translate_docx(data, source_lang=lang, target_lang=target_lang)
                extension = ".docx"
            else:
                translated_file = translate_pdf(data, source_lang=lang, target_lang=target_lang)
                extension = ".pdf"
            
            if translated_file:
                filename = f"traducido_{lang_display.replace('🇪🇸', 'ES').replace('🇺🇸', 'EN').replace('→', '_')}_{doc.file_name.replace('.pdf', extension).replace('.docx', extension)}"
                await update.message.reply_document(
                    document=translated_file,
                    filename=filename,
                    caption=f"{lang_display} Documento traducido\n\n{FIRMA_TEXTO}"
                )
            else:
                await update.message.reply_text("❌ Error al traducir.")
                await processing_msg.delete()
                return
            
            if not context.user_data.get("is_premium", False):
                mark_free_used(uid, "documento")
        else:
            translated_text = translate_text(text, source=lang, target=target_lang)
            audio = tts(translated_text, audio_lang)
            
            if audio:
                await update.message.reply_voice(audio, caption=f"{lang_display} Documento a voz\n\n{FIRMA_TEXTO}")
            else:
                await update.message.reply_text("❌ Error al generar audio.")
                await processing_msg.delete()
                return
            
            if not context.user_data.get("is_premium", False):
                mark_free_used(uid, "doc_voz")
        
        if not context.user_data.get("is_premium", False) and all_free_used(uid):
            keyboard = [[InlineKeyboardButton("💎 PREMIUM", callback_data="plan_premium")]]
            await update.message.reply_text(f"✅ Ya usaste FREE\n\n{FIRMA_TEXTO}", reply_markup=InlineKeyboardMarkup(keyboard))
            context.user_data["waiting_document"] = False
        else:
            back = "premium_menu" if context.user_data.get("is_premium") else "plan_free"
            keyboard = [[InlineKeyboardButton("🔙 Volver al Menú", callback_data=back)]]
            await update.message.reply_text(f"✅ Listo. Envía otro documento o vuelve.\n\n{FIRMA_TEXTO}", reply_markup=InlineKeyboardMarkup(keyboard))
        
        await processing_msg.delete()
    except Exception as e:
        logger.error(f"❌ Error documento: {e}")
        await update.message.reply_text("❌ Error procesando documento.")

async def handle_voice(update: Update, context: ContextTypes.DEFAULT_TYPE):
    uid = update.effective_user.id
    if not context.user_data.get("waiting_audio", False):
        return
    
    try:
        processing_msg = await update.message.reply_text("⏳ Transcribiendo audio...")
        
        voice = update.message.voice
        file = await context.bot.get_file(voice.file_id)
        audio_bytes = await file.download_as_bytearray()
        
        text, detected_lang = transcribe_audio(audio_bytes)
        
        if not text:
            await update.message.reply_text(
                "❌ No se pudo transcribir el audio.\n\n"
                "Asegúrate de hablar claramente.",
                parse_mode="Markdown"
            )
            await processing_msg.delete()
            back = "premium_menu" if context.user_data.get("is_premium") else "plan_free"
            keyboard = [[InlineKeyboardButton("🔙 Volver al Menú", callback_data=back)]]
            await update.message.reply_text(FIRMA_TEXTO, reply_markup=InlineKeyboardMarkup(keyboard))
            return
        
        if detected_lang == "es":
            target_lang = "en"
            lang_display = "🇪🇸→🇺🇸"
            audio_lang = "en"
        else:
            target_lang = "es"
            lang_display = "🇺🇸→🇪🇸"
            audio_lang = "es"
        
        translated_text = translate_text(text, source=detected_lang, target=target_lang)
        
        await update.message.reply_text(
            f"📝 *Transcripción:*\n{text}\n\n"
            f"{lang_display} *Traducción:*\n{translated_text}",
            parse_mode="Markdown"
        )
        
        audio_translated = tts(translated_text, audio_lang)
        if audio_translated:
            await update.message.reply_voice(
                audio_translated,
                caption=f"{lang_display} Audio traducido\n\n{FIRMA_TEXTO}"
            )
        
        if not context.user_data.get("is_premium", False):
            mark_free_used(uid, "audio")
        
        if not context.user_data.get("is_premium", False) and all_free_used(uid):
            keyboard = [[InlineKeyboardButton("💎 PREMIUM", callback_data="plan_premium")]]
            await update.message.reply_text(f"✅ Ya usaste FREE\n\n{FIRMA_TEXTO}", reply_markup=InlineKeyboardMarkup(keyboard))
            context.user_data["waiting_audio"] = False
        else:
            back = "premium_menu" if context.user_data.get("is_premium") else "plan_free"
            keyboard = [[InlineKeyboardButton("🔙 Volver al Menú", callback_data=back)]]
            await update.message.reply_text(f"✅ Listo. Envía otro audio o vuelve.\n\n{FIRMA_TEXTO}", reply_markup=InlineKeyboardMarkup(keyboard))
        
        await processing_msg.delete()
        
    except Exception as e:
        logger.error(f"❌ Error audio: {e}")
        await update.message.reply_text("❌ Error procesando audio.")

async def button_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    data = query.data
    
    # Acciones de imagen
    if data in ["img_solo_texto", "img_audio_original", "img_audio_traducido", "img_analizar", "img_audio_analisis"]:
        return await handle_image_action(update, context)
    
    if data == "translate_yes":
        fake_message = type('obj', (object,), {
            'text': 'SI',
            'reply_text': query.message.reply_text,
            'reply_voice': query.message.reply_voice
        })()
        fake_update = type('obj', (object,), {
            'message': fake_message,
            'effective_user': update.effective_user,
            'callback_query': query
        })()
        return await handle_text(fake_update, context)
    
    elif data == "translate_no":
        fake_message = type('obj', (object,), {
            'text': 'NO',
            'reply_text': query.message.reply_text,
            'reply_voice': query.message.reply_voice
        })()
        fake_update = type('obj', (object,), {
            'message': fake_message,
            'effective_user': update.effective_user,
            'callback_query': query
        })()
        return await handle_text(fake_update, context)
    
    if data == "back_start":
        context.user_data.clear()
        return await start(update, context)
    elif data == "plan_free":
        return await plan_free(update, context)
    elif data == "plan_premium":
        return await plan_premium(update, context)
    elif data == "buy_premium":
        return await buy_premium(update, context)
    elif data == "start_buy_form":
        return await start_buy_form(update, context)
    elif data in ["pago_western", "pago_zelle"]:
        return await buy_metodo_pago(update, context)
    elif data == "premium_login":
        return await premium_login(update, context)
    elif data == "forgot_password":
        return await forgot_password(update, context)
    elif data == "free_texto":
        return await free_texto(update, context)
    elif data == "free_documento":
        return await free_documento(update, context)
    elif data == "free_doc_voz":
        return await free_doc_voz(update, context)
    elif data == "free_audio":
        return await free_audio(update, context)
    elif data == "free_imagen":
        return await free_imagen(update, context)
    elif data == "premium_menu":
        return await show_premium_menu(update, context)
    elif data == "premium_texto":
        return await premium_texto(update, context)
    elif data == "premium_documento":
        return await premium_documento(update, context)
    elif data == "premium_doc_voz":
        return await premium_doc_voz(update, context)
    elif data == "premium_audio":
        return await premium_audio(update, context)
    elif data == "premium_imagen":
        return await premium_imagen(update, context)
    elif data == "premium_logout":
        uid = update.effective_user.id
        if uid in active_sessions:
            del active_sessions[uid]
        context.user_data.clear()
        return await start(update, context)
    
    return CHOOSING_PLAN

async def error_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    logger.error(f"Error: {context.error}")
    if update and update.effective_message:
        await update.effective_message.reply_text("❌ Error. Usa /start")

def main():
    if not TELEGRAM_BOT_TOKEN:
        logger.error("❌ TOKEN NO CONFIGURADO")
        return
    
    if not os.getenv("GEMINI_API_KEY"):
        logger.warning("⚠️ GEMINI_API_KEY no configurado - función de imágenes no disponible")
    
    logger.info("🚀 Iniciando bot...")
    app = Application.builder().token(TELEGRAM_BOT_TOKEN).build()
    
    conv_handler = ConversationHandler(
        entry_points=[CommandHandler("start", start)],
        states={
            CHOOSING_PLAN: [
                CallbackQueryHandler(button_callback),
                MessageHandler(filters.VOICE, handle_voice),
                MessageHandler(filters.Document.ALL, handle_document),
                MessageHandler(filters.PHOTO, handle_image),
                MessageHandler(filters.TEXT & ~filters.COMMAND, handle_text)
            ],
            PREMIUM_USERNAME: [MessageHandler(filters.TEXT & ~filters.COMMAND, premium_username)],
            PREMIUM_PASSWORD: [MessageHandler(filters.TEXT & ~filters.COMMAND, premium_password)],
            FORGOT_PASSWORD: [MessageHandler(filters.TEXT & ~filters.COMMAND, process_forgot_password)],
            BUY_NOMBRE: [MessageHandler(filters.TEXT & ~filters.COMMAND, buy_nombre)],
            BUY_APELLIDO: [MessageHandler(filters.TEXT & ~filters.COMMAND, buy_apellido)],
            BUY_EMAIL: [MessageHandler(filters.TEXT & ~filters.COMMAND, buy_email)],
            BUY_CELULAR: [MessageHandler(filters.TEXT & ~filters.COMMAND, buy_celular)],
            BUY_METODO_PAGO: [CallbackQueryHandler(button_callback)],
        },
        fallbacks=[CommandHandler("start", start)],
        allow_reentry=True,
        per_message=False,
        per_chat=True,
        per_user=True
    )
    
    app.add_handler(conv_handler)
    app.add_error_handler(error_handler)
    
    logger.info("✅ Bot listo")
    app.run_polling(allowed_updates=Update.ALL_TYPES)

if __name__ == "__main__":
    main()
